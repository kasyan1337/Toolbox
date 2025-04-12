from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os


def clean_up_docx(input_file, output_folder):
    """
    Cleans up a DOCX file by removing background colors, setting text color to black,
    and standardizing the font. Saves the cleaned DOCX to the specified output folder.

    :param input_file: Path to the DOCX file to clean up
    :param output_folder: Path to the folder where cleaned DOCX will be saved
    """
    # Load the DOCX document
    doc = Document(input_file)

    # Iterate over each paragraph to remove background colors and standardize font
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.highlight_color = None  # Remove highlight/background color
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
            run.font.name = "Arial"  # Set a default font, e.g., Arial
            run.font.size = None  # Reset font size if necessary

    # Remove background colors from table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Set the background color of the cell to white
                cell._element.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls("w")))
                )

    # Prepare the output path
    output_file = os.path.join(output_folder, f"cleaned_{os.path.basename(input_file)}")

    # Save the cleaned DOCX document
    doc.save(output_file)
    print(f"Cleaned document saved as: {output_file}")


if __name__ == "__main__":
    # Example usage
    input_file = "output/test_converted.docx"  # Path to the file you want to clean up
    output_folder = "output"  # Folder to save the cleaned file

    # Ensure output directory exists
    os.makedirs(output_folder, exist_ok=True)

    # Run the cleanup function
    clean_up_docx(input_file, output_folder)
